"""Shared test fixtures."""

import sys
import os
import tempfile
import shutil
from pathlib import Path
import pytest

# Add src to path
sys.path.insert(0, '/app/src')
os.environ.setdefault('OLLAMA_HOST', 'http://localhost:11434')


@pytest.fixture
def tmp_dir():
    """Provide a temporary directory that's cleaned up after the test."""
    d = Path(tempfile.mkdtemp())
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture
def sample_config():
    """Return a minimal processing config dict."""
    return {
        'supported_formats': ['pdf', 'docx', 'pptx', 'html', 'csv', 'xlsx', 'png', 'jpg'],
        'ocr': {'enabled': False, 'language': 'eng', 'min_confidence': 60},
        'images': {'extract': True, 'max_size_mb': 10},
        'ai': {'generate_summary': True, 'generate_tags': True, 'max_tags': 10},
        'csv': {'max_rows': 100},
        'xlsx': {'max_rows': 100, 'include_formulas': False},
        'docling': {'table_mode': 'fast', 'large_pdf_page_threshold': 30, 'xlarge_pdf_page_threshold': 80},
    }


@pytest.fixture
def sample_obsidian_config():
    """Return a minimal obsidian config dict."""
    return {
        'frontmatter': {
            'add_source_file': True,
            'add_created_date': True,
            'add_processed_date': True,
            'add_ai_summary': True,
            'add_ai_tags': True,
        },
        'preserve_filename': True,
        'sanitize_filenames': True,
        'attachments_folder': 'attachments',
        'embed_images': True,
    }


@pytest.fixture
def make_pdf(tmp_dir):
    """Create a minimal valid PDF file."""
    def _make(name='test.pdf', pages=1):
        path = tmp_dir / name
        # Minimal valid PDF
        content = b"""%PDF-1.0
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>stream
BT /F1 24 Tf 100 700 Td (Hello World) Tj ET
endstream endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000340 00000 n
trailer<</Size 6/Root 1 0 R>>
startxref
434
%%EOF"""
        path.write_bytes(content)
        return path
    return _make


@pytest.fixture
def make_csv(tmp_dir):
    """Create a CSV test file."""
    def _make(name='test.csv', rows=50):
        path = tmp_dir / name
        import csv
        with open(path, 'w', newline='') as f:
            w = csv.writer(f)
            w.writerow(['id', 'name', 'value'])
            for i in range(rows):
                w.writerow([i, f'item_{i}', i * 1.5])
        return path
    return _make


@pytest.fixture
def make_image(tmp_dir):
    """Create a test image file."""
    def _make(name='test.png', width=100, height=100, fmt='PNG'):
        from PIL import Image
        path = tmp_dir / name
        img = Image.new('RGB', (width, height), color=(255, 0, 0))
        img.save(path, fmt)
        return path
    return _make


@pytest.fixture
def make_xlsx(tmp_dir):
    """Create a test XLSX file."""
    def _make(name='test.xlsx', rows=20):
        import openpyxl
        path = tmp_dir / name
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'TestSheet'
        ws.append(['Name', 'Age', 'Score'])
        for i in range(rows):
            ws.append([f'Person_{i}', 20 + i, 50 + i * 2])
        wb.save(path)
        return path
    return _make


@pytest.fixture
def make_docx(tmp_dir):
    """Create a test DOCX file."""
    def _make(name='test.docx'):
        from docx import Document
        path = tmp_dir / name
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph with some content.')
        doc.add_heading('Section One', level=1)
        doc.add_paragraph('Content in section one.')
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = 'Header A'
        table.cell(0, 1).text = 'Header B'
        table.cell(1, 0).text = 'Row 1A'
        table.cell(1, 1).text = 'Row 1B'
        table.cell(2, 0).text = 'Row 2A'
        table.cell(2, 1).text = 'Row 2B'
        doc.save(path)
        return path
    return _make
