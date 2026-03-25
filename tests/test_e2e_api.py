"""End-to-end API tests against the live running container.

Tests the /api/convert sync endpoint and /api/upload async flow
with real file conversions (no mocks).
"""

import time
import tempfile
import csv
import requests
from pathlib import Path

BASE_URL = "http://localhost:3143"
TIMEOUT = 180  # seconds


def _create_csv(path: Path, rows=10):
    with open(path, 'w', newline='') as f:
        w = csv.writer(f)
        w.writerow(['id', 'name', 'value'])
        for i in range(rows):
            w.writerow([i, f'item_{i}', i * 1.5])


def _create_pdf(path: Path):
    path.write_bytes(b"""%PDF-1.0
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>stream
BT /F1 24 Tf 100 700 Td (Hello World) Tj ET
endstream endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000340 00000 n
trailer<</Size 6/Root 1 0 R>>
startxref
434
%%EOF""")


def _create_image(path: Path):
    from PIL import Image
    img = Image.new('RGB', (100, 100), color=(255, 0, 0))
    img.save(path, 'PNG')


def _create_docx(path: Path):
    from docx import Document
    doc = Document()
    doc.add_heading('API Test Document', 0)
    doc.add_paragraph('Content for API end-to-end testing.')
    doc.save(path)


def _create_xlsx(path: Path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Data'
    ws.append(['Name', 'Score'])
    for i in range(5):
        ws.append([f'Person_{i}', 80 + i])
    wb.save(path)


# ─── Health & Meta ───────────────────────────────────────────────

class TestHealthAPI:
    def test_health(self):
        r = requests.get(f"{BASE_URL}/health")
        assert r.status_code == 200
        assert r.json()['status'] == 'healthy'

    def test_stats_removed(self):
        """Stats endpoint removed for multi-user privacy."""
        r = requests.get(f"{BASE_URL}/api/stats")
        assert r.status_code in [404, 405]

    def test_no_cache_headers(self):
        r = requests.get(f"{BASE_URL}/api/jobs/test-id")
        assert 'no-store' in r.headers.get('Cache-Control', '')


# ─── Sync /api/convert ──────────────────────────────────────────

class TestSyncConvertAPI:
    def test_no_file_422(self):
        """FastAPI returns 422 when required 'file' field is missing."""
        r = requests.post(f"{BASE_URL}/api/convert")
        assert r.status_code == 422

    def test_unsupported_format_415(self):
        r = requests.post(f"{BASE_URL}/api/convert",
                          files={'file': ('test.mp4', b'\x00\x00', 'video/mp4')})
        assert r.status_code == 415
        d = r.json()
        detail = d.get('detail', d)
        assert 'Unsupported' in detail.get('error', '')
        assert 'supported_formats' in detail

    def test_unsupported_mp3_415(self):
        r = requests.post(f"{BASE_URL}/api/convert",
                          files={'file': ('test.mp3', b'\xff\xfb', 'audio/mpeg')})
        assert r.status_code == 415

    def test_unsupported_exe_415(self):
        r = requests.post(f"{BASE_URL}/api/convert",
                          files={'file': ('test.exe', b'MZ', 'application/octet-stream')})
        assert r.status_code == 415

    def test_csv_json_output(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'api_test.csv'
            _create_csv(p, rows=5)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            d = r.json()
            assert d['status'] == 'completed'
            assert d['filename'] == 'api_test.csv'
            assert '5 rows' in d['markdown']
            assert 'item_0' in d['markdown']

    def test_csv_markdown_output(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'md_test.csv'
            _create_csv(p, rows=3)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'markdown', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            assert r.headers['Content-Type'].startswith('text/markdown')
            assert 'item_0' in r.text

    def test_csv_truncation(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'big.csv'
            _create_csv(p, rows=2000)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            assert 'Showing first' in r.json()['markdown']

    def test_image_conversion(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'api_test.png'
            _create_image(p)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            d = r.json()
            assert d['status'] == 'completed'
            assert '100' in d['markdown']  # dimensions
            assert 'PNG' in d['markdown']

    def test_docx_conversion(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'api_test.docx'
            _create_docx(p)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            d = r.json()
            assert d['status'] == 'completed'
            assert 'API Test Document' in d['markdown']
            assert 'source_type: docx' in d['markdown']

    def test_xlsx_conversion(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'api_test.xlsx'
            _create_xlsx(p)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            d = r.json()
            assert d['status'] == 'completed'
            assert 'Person_0' in d['markdown']
            assert 'sheet_count' in d['markdown']

    def test_html_entities_unescaped(self):
        """Verify &amp; doesn't appear in output (Docling HTML entity fix)."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'entities.docx'
            from docx import Document
            doc = Document()
            doc.add_paragraph('Sales & Marketing R&D')
            doc.save(p)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            md = r.json()['markdown']
            assert '&amp;' not in md
            assert '&' in md  # raw ampersand should be present

    def test_source_type_lowercase(self):
        """Verify source_type is always lowercase."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'api_test.png'
            _create_image(p)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': ('TEST.PNG', open(p, 'rb'), 'image/png')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            assert 'source_type: png' in r.json()['markdown']

    def test_ai_enhancement_on(self):
        """With AI enabled, frontmatter should have summary and tags."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'ai_test.csv'
            _create_csv(p, rows=20)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'true'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            md = r.json()['markdown']
            assert 'summary:' in md
            assert 'tags:' in md
            assert 'description:' in md

    def test_ai_enhancement_off(self):
        """With AI disabled, no summary/tags in frontmatter."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'noai_test.csv'
            _create_csv(p, rows=5)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            md = r.json()['markdown']
            assert 'summary:' not in md.split('---')[1]  # not in frontmatter

    def test_cleanup_no_temp_files(self):
        """After conversion, no temp files should remain."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'cleanup_test.csv'
            _create_csv(p, rows=3)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            # Check no leftover in _api_tmp or output
            import subprocess
            result = subprocess.run(
                ['docker', 'exec', 'myfiles_web', 'ls', '/app/uploads/_api_tmp/'],
                capture_output=True, text=True
            )
            assert 'cleanup_test.csv' not in result.stdout


# ─── Async /api/upload flow ──────────────────────────────────────

class TestAsyncUploadAPI:
    def test_upload_and_poll(self):
        """Upload via async, poll for completion, then download."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'async_test.csv'
            _create_csv(p, rows=5)

            # Upload
            r = requests.post(f"{BASE_URL}/api/upload",
                              files={'files[]': open(p, 'rb')},
                              data={'ai_enhancement': 'false'})
            assert r.status_code == 200
            job_id = r.json()['job_ids'][0]

            # Poll until complete
            for _ in range(60):
                r = requests.get(f"{BASE_URL}/api/jobs/{job_id}")
                assert r.status_code == 200
                status = r.json()['status']
                if status == 'completed':
                    break
                if status == 'failed':
                    raise AssertionError(f"Job failed: {r.json().get('error')}")
                time.sleep(1)
            else:
                raise AssertionError("Job did not complete within 60 seconds")

            # Download
            r = requests.get(f"{BASE_URL}/api/download/{job_id}")
            assert r.status_code == 200
            assert 'item_0' in r.text

    def test_upload_unsupported_skipped(self):
        r = requests.post(f"{BASE_URL}/api/upload",
                          files={'files[]': ('bad.mp4', b'\x00', 'video/mp4')},
                          data={'ai_enhancement': 'false'})
        d = r.json()
        assert d['job_ids'] == []
        assert 'bad.mp4' in d.get('skipped', [])

    def test_upload_mixed_files(self):
        """Upload one supported and one unsupported file."""
        with tempfile.TemporaryDirectory() as tmp:
            good = Path(tmp) / 'good.csv'
            _create_csv(good, rows=3)
            r = requests.post(f"{BASE_URL}/api/upload",
                              files=[
                                  ('files[]', ('good.csv', open(good, 'rb'), 'text/csv')),
                                  ('files[]', ('bad.exe', b'MZ', 'application/octet-stream')),
                              ],
                              data={'ai_enhancement': 'false'})
            d = r.json()
            assert len(d['job_ids']) == 1
            assert 'bad.exe' in d.get('skipped', [])

    def test_concurrent_same_filename(self):
        """Two users uploading same filename get correct independent results."""
        import concurrent.futures as cf

        def upload(uid):
            with tempfile.TemporaryDirectory() as tmp:
                p = Path(tmp) / 'report.csv'
                with open(p, 'w', newline='') as f:
                    w = csv.writer(f)
                    w.writerow(['user', 'val'])
                    w.writerow([f'user_{uid}', uid * 99])
                r = requests.post(f"{BASE_URL}/api/convert",
                                  files={'file': open(p, 'rb')},
                                  data={'output_format': 'json', 'ai_enhancement': 'false'},
                                  timeout=60)
                return uid, r.status_code, r.json()

        with cf.ThreadPoolExecutor(max_workers=3) as pool:
            results = list(pool.map(lambda uid: upload(uid), range(3)))

        for uid, code, data in results:
            assert code == 200, f"User {uid} got HTTP {code}"
            assert data['status'] == 'completed'
            assert f'user_{uid}' in data['markdown'], f"User {uid} got wrong data"

    def test_download_nonexistent(self):
        r = requests.get(f"{BASE_URL}/api/download/nonexistent-id")
        assert r.status_code == 404

    def test_job_nonexistent(self):
        r = requests.get(f"{BASE_URL}/api/jobs/nonexistent-id")
        assert r.status_code == 404


# ─── PDF via sync (real Docling conversion) ──────────────────────

class TestPdfConversion:
    def test_minimal_pdf(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'hello.pdf'
            _create_pdf(p)
            r = requests.post(f"{BASE_URL}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=TIMEOUT)
            assert r.status_code == 200
            d = r.json()
            assert d['status'] == 'completed'
            assert 'source_type: pdf' in d['markdown']


# ─── Chunked Upload ──────────────────────────────────────────────

class TestChunkedUploadAPI:
    def test_full_chunked_flow(self):
        """Init → chunk → complete → poll → download."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'chunked.csv'
            _create_csv(p, rows=10)
            data = p.read_bytes()

            # Init
            r = requests.post(f"{BASE_URL}/api/upload/init", data={
                'filename': 'chunked.csv',
                'total_size': len(data),
                'total_chunks': 1,
                'ai_enhancement': 'false',
            })
            assert r.status_code == 200
            upload_id = r.json()['upload_id']

            # Chunk
            r = requests.post(f"{BASE_URL}/api/upload/chunk/{upload_id}",
                              files={'chunk': ('chunk_0', data)},
                              data={'chunk_index': 0})
            assert r.status_code == 200
            assert r.json()['received'] == 1

            # Complete
            r = requests.post(f"{BASE_URL}/api/upload/complete/{upload_id}")
            assert r.status_code == 200
            job_id = r.json()['job_id']

            # Poll
            for _ in range(30):
                r = requests.get(f"{BASE_URL}/api/jobs/{job_id}")
                if r.json()['status'] == 'completed':
                    break
                time.sleep(1)
            assert r.json()['status'] == 'completed'

            # Download
            r = requests.get(f"{BASE_URL}/api/download/{job_id}")
            assert r.status_code == 200
            assert 'item_0' in r.text

    def test_multi_chunk_flow(self):
        """Upload a file in 2 chunks."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'multi_chunk.csv'
            _create_csv(p, rows=50)
            data = p.read_bytes()
            mid = len(data) // 2

            # Init
            r = requests.post(f"{BASE_URL}/api/upload/init", data={
                'filename': 'multi_chunk.csv',
                'total_size': len(data),
                'total_chunks': 2,
                'ai_enhancement': 'false',
            })
            upload_id = r.json()['upload_id']

            # Chunk 0
            r = requests.post(f"{BASE_URL}/api/upload/chunk/{upload_id}",
                              files={'chunk': ('c0', data[:mid])},
                              data={'chunk_index': 0})
            assert r.json()['received'] == 1

            # Chunk 1
            r = requests.post(f"{BASE_URL}/api/upload/chunk/{upload_id}",
                              files={'chunk': ('c1', data[mid:])},
                              data={'chunk_index': 1})
            assert r.json()['received'] == 2

            # Complete
            r = requests.post(f"{BASE_URL}/api/upload/complete/{upload_id}")
            assert r.status_code == 200
            job_id = r.json()['job_id']

            # Poll
            for _ in range(30):
                r = requests.get(f"{BASE_URL}/api/jobs/{job_id}")
                if r.json()['status'] == 'completed':
                    break
                time.sleep(1)
            assert r.json()['status'] == 'completed'

    def test_init_unsupported_format(self):
        r = requests.post(f"{BASE_URL}/api/upload/init", data={
            'filename': 'video.mp4',
            'total_size': 1000,
            'total_chunks': 1,
        })
        assert r.status_code == 415

    def test_chunk_invalid_session(self):
        r = requests.post(f"{BASE_URL}/api/upload/chunk/nonexistent-id",
                          files={'chunk': ('c', b'data')},
                          data={'chunk_index': 0})
        assert r.status_code == 404

    def test_complete_missing_chunks(self):
        """Complete before all chunks uploaded should fail."""
        r = requests.post(f"{BASE_URL}/api/upload/init", data={
            'filename': 'incomplete.csv',
            'total_size': 1000,
            'total_chunks': 5,
            'ai_enhancement': 'false',
        })
        upload_id = r.json()['upload_id']

        # Only send 1 of 5 chunks
        requests.post(f"{BASE_URL}/api/upload/chunk/{upload_id}",
                      files={'chunk': ('c', b'partial')},
                      data={'chunk_index': 0})

        r = requests.post(f"{BASE_URL}/api/upload/complete/{upload_id}")
        assert r.status_code == 400
        assert 'Missing' in r.json().get('detail', {}).get('error', '')


# ─── Swagger / OpenAPI ───────────────────────────────────────────

class TestOpenAPI:
    def test_swagger_ui(self):
        r = requests.get(f"{BASE_URL}/docs")
        assert r.status_code == 200
        assert 'swagger' in r.text.lower() or 'openapi' in r.text.lower()

    def test_openapi_json(self):
        r = requests.get(f"{BASE_URL}/openapi.json")
        assert r.status_code == 200
        d = r.json()
        assert d['info']['title'] == 'MyFiles to Markdown API'
        assert '/api/convert' in d['paths']
        assert '/api/upload/init' in d['paths']
        assert '/api/upload/chunk/{upload_id}' in d['paths']
        assert '/api/upload/complete/{upload_id}' in d['paths']
        # Stats endpoint removed for multi-user privacy
        assert '/api/stats' not in d['paths']


if __name__ == '__main__':
    import pytest
    pytest.main([__file__, '-v', '--tb=short'])
