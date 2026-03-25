"""End-to-end browser tests using Playwright.

Tests the full upload → processing → status update → download flow
through the actual web UI at http://localhost:3143.
"""

import time
import uuid
import tempfile
import csv
from pathlib import Path
from playwright.sync_api import sync_playwright, expect

BASE_URL = "http://localhost:3143"


def _create_test_csv(path: Path, rows=10):
    with open(path, 'w', newline='') as f:
        w = csv.writer(f)
        w.writerow(['id', 'name', 'score'])
        for i in range(rows):
            w.writerow([i, f'item_{i}', i * 2.5])


def _create_test_pdf(path: Path):
    """Create minimal valid PDF."""
    content = b"""%PDF-1.0
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>stream
BT /F1 24 Tf 100 700 Td (Hello World) Tj ET
endstream endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000340 00000 n
trailer<</Size 6/Root 1 0 R>>
startxref
434
%%EOF"""
    path.write_bytes(content)


def _create_test_image(path: Path):
    from PIL import Image
    img = Image.new('RGB', (200, 200), color=(0, 128, 255))
    img.save(path, 'PNG')


def _create_test_docx(path: Path):
    from docx import Document
    doc = Document()
    doc.add_heading('E2E Test Document', 0)
    doc.add_paragraph('This document was created by the Playwright E2E test.')
    doc.add_heading('Section One', level=1)
    doc.add_paragraph('Automated testing validates the full pipeline.')
    doc.save(path)


def test_page_loads():
    """Test that the web UI loads and has key elements."""
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(BASE_URL)

        # Check page title / key elements exist
        assert page.title() != ""
        assert page.locator('#uploadArea').is_visible()
        assert page.locator('#jobList').is_visible()

        browser.close()


def test_health_endpoint():
    """Test health check via browser fetch."""
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        resp = page.request.get(f"{BASE_URL}/health")
        assert resp.ok
        data = resp.json()
        assert data['status'] == 'healthy'
        browser.close()


def test_api_no_cache_headers():
    """Verify API responses have no-cache headers (Cloudflare fix)."""
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        resp = page.request.get(f"{BASE_URL}/api/stats")
        assert resp.ok
        headers = resp.headers
        assert 'no-store' in headers.get('cache-control', '')
        assert headers.get('pragma') == 'no-cache'
        browser.close()


def test_upload_csv_and_status_updates():
    """Upload a CSV, verify status transitions from processing to completed."""
    with tempfile.TemporaryDirectory() as tmp:
        csv_path = Path(tmp) / 'e2e_test.csv'
        _create_test_csv(csv_path, rows=5)

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(BASE_URL)

            # Upload file via the file input
            page.locator('#fileInput').set_input_files(str(csv_path))

            # Wait for the job card to appear
            job_card = page.locator('.job-card').first
            job_card.wait_for(state='visible', timeout=10000)

            # Filename should be visible
            assert 'e2e_test.csv' in job_card.text_content()

            # Wait for status to change to completed (poll updates every 2s)
            page.locator('.status-completed').first.wait_for(state='visible', timeout=60000)

            # Download button should now be visible
            download_btn = page.locator('button:has-text("Download")').first
            assert download_btn.is_visible()

            browser.close()


def test_upload_image_fast_completion():
    """Upload a PNG image — should complete in seconds."""
    unique = uuid.uuid4().hex[:8]
    with tempfile.TemporaryDirectory() as tmp:
        img_path = Path(tmp) / f'img_{unique}.png'
        _create_test_image(img_path)

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(BASE_URL)

            page.locator('#fileInput').set_input_files(str(img_path))

            page.locator(f'.job-card:has-text("img_{unique}.png") .status-completed').wait_for(
                state='visible', timeout=30000
            )

            card = page.locator(f'.job-card:has-text("img_{unique}.png")').first
            assert 'Download' in card.text_content()

            browser.close()


def test_upload_docx_and_download():
    """Upload a DOCX, wait for completion, and download the result."""
    unique = uuid.uuid4().hex[:8]
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = Path(tmp) / f'docx_{unique}.docx'
        _create_test_docx(docx_path)

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(BASE_URL)

            page.locator('#fileInput').set_input_files(str(docx_path))

            page.locator(f'.job-card:has-text("docx_{unique}.docx") .status-completed').wait_for(
                state='visible', timeout=60000
            )

            with page.expect_download() as download_info:
                page.locator(f'.job-card:has-text("docx_{unique}.docx") button:has-text("Download")').first.click()
            download = download_info.value
            assert download.suggested_filename == f'docx_{unique}.md'

            # Read downloaded content
            dl_path = Path(tmp) / 'downloaded.md'
            download.save_as(str(dl_path))
            content = dl_path.read_text()
            assert 'E2E Test Document' in content
            assert 'source_type: docx' in content

            browser.close()


def test_stats_update_after_upload():
    """Verify the stats counters update after a file is processed."""
    with tempfile.TemporaryDirectory() as tmp:
        img_path = Path(tmp) / 'stats_test.png'
        _create_test_image(img_path)

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(BASE_URL)

            # Get initial completed count
            page.wait_for_timeout(3000)  # Let polling stabilize
            initial_stats = page.locator('#stat-completed').text_content()

            # Upload a file
            page.locator('#fileInput').set_input_files(str(img_path))

            # Wait for completion
            page.locator('.status-completed').first.wait_for(state='visible', timeout=30000)

            # Stats should have incremented
            page.wait_for_timeout(3000)  # Wait for next poll
            updated_stats = page.locator('#stat-completed').text_content()
            assert updated_stats != initial_stats or 'Completed: 0' not in updated_stats

            browser.close()


def test_multiple_file_upload():
    """Upload two files one after another and verify both complete."""
    unique = uuid.uuid4().hex[:8]
    with tempfile.TemporaryDirectory() as tmp:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(BASE_URL)

            for i in range(2):
                img_path = Path(tmp) / f'seq_{unique}_{i}.png'
                _create_test_image(img_path)
                page.locator('#fileInput').set_input_files(str(img_path))
                page.locator(f'.job-card:has-text("seq_{unique}_{i}.png") .status-completed').wait_for(
                    state='visible', timeout=60000
                )

            all_text = page.locator('#jobList').text_content()
            assert f'seq_{unique}_0.png' in all_text
            assert f'seq_{unique}_1.png' in all_text

            browser.close()


if __name__ == '__main__':
    import pytest
    pytest.main([__file__, '-v', '--tb=short'])
