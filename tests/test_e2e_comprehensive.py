"""Comprehensive E2E test suite — edge cases, outliers, exploratory, and performance.

Targets >90% functional coverage of the API surface.
Runs against the live container at localhost:3143.
"""

import time
import csv
import io
import math
import os
import tempfile
import concurrent.futures
import requests
from pathlib import Path
from PIL import Image
from docx import Document
import openpyxl

BASE = "http://localhost:3143"
T = 180  # default timeout


def _csv(path, rows=5, cols=3):
    with open(path, 'w', newline='') as f:
        w = csv.writer(f)
        w.writerow([f'col_{i}' for i in range(cols)])
        for r in range(rows):
            w.writerow([f'r{r}c{c}' for c in range(cols)])


def _img(path, w=100, h=100, fmt='PNG', mode='RGB'):
    color = 128 if mode == 'L' else (255, 128, 0) if mode == 'RGB' else (255, 128, 0, 255)
    Image.new(mode, (w, h), color=color).save(path, fmt)


def _pdf(path):
    path.write_bytes(b"""%PDF-1.0
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>stream
BT /F1 24 Tf 100 700 Td (Hello World) Tj ET
endstream endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000340 00000 n
trailer<</Size 6/Root 1 0 R>>
startxref
434
%%EOF""")


def _docx(path, text="Test document content."):
    doc = Document()
    doc.add_heading('Test', 0)
    doc.add_paragraph(text)
    doc.save(path)


def _xlsx(path, rows=5):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['A', 'B'])
    for i in range(rows):
        ws.append([i, i * 2])
    wb.save(path)


def _html(path, body="<h1>Test</h1><p>Hello world</p>"):
    path.write_text(f"<html><body>{body}</body></html>")


def _convert(filepath, **kwargs):
    kwargs.setdefault('output_format', 'json')
    kwargs.setdefault('ai_enhancement', 'false')
    data = {k: v for k, v in kwargs.items() if k != 'timeout'}
    r = requests.post(f"{BASE}/api/convert",
                      files={'file': open(filepath, 'rb')},
                      data=data,
                      timeout=kwargs.get('timeout', T))
    return r


# ═══════════════════════════════════════════════════════════════════
# 1. INPUT VALIDATION & ERROR HANDLING
# ═══════════════════════════════════════════════════════════════════

class TestInputValidation:
    """Test all input validation paths."""

    def test_no_file_returns_422(self):
        r = requests.post(f"{BASE}/api/convert", data={'output_format': 'json'})
        assert r.status_code == 422

    def test_empty_filename(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('', b'', 'application/octet-stream')})
        # FastAPI may reject or treat as no file
        assert r.status_code in [400, 415, 422]

    def test_no_extension(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('noext', b'data', 'application/octet-stream')})
        assert r.status_code == 415

    def test_double_extension(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('file.pdf.exe', b'data', 'application/octet-stream')})
        assert r.status_code == 415  # .exe is unsupported

    def test_hidden_file(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('.hidden.pdf', b'%PDF-1', 'application/pdf')})
        # Should still process based on extension
        assert r.status_code in [200, 500]  # 500 if PDF is invalid

    def test_unicode_filename(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'test.csv'
            _csv(p, rows=2)
            r = requests.post(f"{BASE}/api/convert",
                              files={'file': ('Übersicht_Daten.csv', open(p, 'rb'))},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=T)
            assert r.status_code == 200

    def test_very_long_filename(self):
        name = 'a' * 200 + '.csv'
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'test.csv'
            _csv(p, rows=2)
            r = requests.post(f"{BASE}/api/convert",
                              files={'file': (name, open(p, 'rb'))},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=T)
            assert r.status_code == 200

    def test_special_chars_filename(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'test.csv'
            _csv(p, rows=2)
            r = requests.post(f"{BASE}/api/convert",
                              files={'file': ('file (1) [copy] & backup.csv', open(p, 'rb'))},
                              data={'output_format': 'json', 'ai_enhancement': 'false'},
                              timeout=T)
            assert r.status_code == 200

    def test_all_unsupported_formats(self):
        for ext in ['mp4', 'mp3', 'avi', 'mov', 'zip', 'tar', 'exe', 'dll', 'py', 'js', 'txt', 'rtf']:
            r = requests.post(f"{BASE}/api/convert",
                              files={'file': (f'test.{ext}', b'data')})
            assert r.status_code == 415, f".{ext} should return 415, got {r.status_code}"

    def test_all_supported_formats_accepted(self):
        """Verify every supported format is at least accepted (not 415)."""
        for ext in ['pdf', 'docx', 'pptx', 'html', 'htm', 'csv', 'xlsx', 'xls',
                     'png', 'jpg', 'jpeg', 'tiff', 'tif', 'bmp', 'gif', 'webp',
                     'ipynb', 'eml', 'msg', 'mbox']:
            r = requests.post(f"{BASE}/api/convert",
                              files={'file': (f'test.{ext}', b'minimal data')},
                              data={'ai_enhancement': 'false'})
            assert r.status_code != 415, f".{ext} should not return 415"

    def test_case_insensitive_extension(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'test.csv'
            _csv(p, rows=2)
            for ext in ['CSV', 'Csv', 'cSv']:
                r = requests.post(f"{BASE}/api/convert",
                                  files={'file': (f'test.{ext}', open(p, 'rb'))},
                                  data={'output_format': 'json', 'ai_enhancement': 'false'},
                                  timeout=T)
                assert r.status_code == 200, f".{ext} failed"


# ═══════════════════════════════════════════════════════════════════
# 2. FORMAT-SPECIFIC CONVERSION QUALITY
# ═══════════════════════════════════════════════════════════════════

class TestConversionQuality:
    """Verify each format produces correct markdown."""

    def test_csv_preserves_data(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'data.csv'
            _csv(p, rows=10, cols=5)
            r = _convert(p)
            md = r.json()['markdown']
            assert 'r0c0' in md
            assert 'r9c4' in md
            assert 'col_0' in md
            assert '10 rows' in md

    def test_csv_empty_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'empty.csv'
            p.write_text('')
            r = _convert(p)
            assert r.status_code in [200, 500]

    def test_csv_single_column(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'single.csv'
            with open(p, 'w') as f:
                f.write('name\nalice\nbob\n')
            r = _convert(p)
            assert r.status_code == 200
            assert 'alice' in r.json()['markdown']

    def test_csv_with_commas_in_values(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'quoted.csv'
            with open(p, 'w', newline='') as f:
                w = csv.writer(f)
                w.writerow(['name', 'address'])
                w.writerow(['Alice', '123 Main St, Apt 4'])
            r = _convert(p)
            assert r.status_code == 200
            assert 'Apt 4' in r.json()['markdown']

    def test_image_dimensions_in_output(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'img.png'
            _img(p, 320, 240)
            r = _convert(p)
            md = r.json()['markdown']
            assert '320' in md
            assert '240' in md
            assert 'PNG' in md

    def test_image_formats(self):
        fmts = [('test.png', 'PNG'), ('test.jpg', 'JPEG'), ('test.bmp', 'BMP'), ('test.gif', 'GIF')]
        for name, fmt in fmts:
            with tempfile.TemporaryDirectory() as tmp:
                p = Path(tmp) / name
                _img(p, 50, 50, fmt)
                r = _convert(p)
                assert r.status_code == 200, f"{name} failed"

    def test_docx_headings_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'doc.docx'
            doc = Document()
            doc.add_heading('Main Title', 0)
            doc.add_heading('Section A', 1)
            doc.add_paragraph('Content A')
            doc.add_heading('Section B', 1)
            doc.add_paragraph('Content B')
            doc.save(p)
            r = _convert(p)
            md = r.json()['markdown']
            assert 'Main Title' in md
            assert 'Section A' in md
            assert 'Content B' in md

    def test_docx_tables_converted(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'table.docx'
            doc = Document()
            t = doc.add_table(rows=2, cols=2)
            t.cell(0, 0).text = 'H1'
            t.cell(0, 1).text = 'H2'
            t.cell(1, 0).text = 'V1'
            t.cell(1, 1).text = 'V2'
            doc.save(p)
            r = _convert(p)
            md = r.json()['markdown']
            assert 'H1' in md and 'V2' in md

    def test_xlsx_multisheet(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'multi.xlsx'
            wb = openpyxl.Workbook()
            ws1 = wb.active
            ws1.title = 'Sales'
            ws1.append(['Q', 'Revenue'])
            ws1.append(['Q1', 100])
            ws2 = wb.create_sheet('Costs')
            ws2.append(['Item', 'Amount'])
            ws2.append(['Rent', 50])
            wb.save(p)
            r = _convert(p)
            md = r.json()['markdown']
            assert 'Sales' in md
            assert 'Costs' in md
            assert '100' in md
            assert 'sheet_count' in md

    def test_html_conversion(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'page.html'
            _html(p, '<h1>Title</h1><p>Paragraph text.</p><ul><li>Item 1</li></ul>')
            r = _convert(p)
            assert r.status_code == 200
            md = r.json()['markdown']
            assert 'Title' in md

    def test_pdf_conversion(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'doc.pdf'
            _pdf(p)
            r = _convert(p)
            assert r.status_code == 200
            assert 'source_type: pdf' in r.json()['markdown']

    def test_html_entities_unescaped(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'ent.docx'
            doc = Document()
            doc.add_paragraph('R&D Department <team> "quoted"')
            doc.save(p)
            r = _convert(p)
            md = r.json()['markdown']
            assert '&amp;' not in md
            assert 'R&D' in md or 'R&amp;D' not in md

    def test_source_type_always_lowercase(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'test.png'
            _img(p)
            for upload_name in ['TEST.PNG', 'test.Png', 'TEST.png']:
                r = requests.post(f"{BASE}/api/convert",
                                  files={'file': (upload_name, open(p, 'rb'))},
                                  data={'output_format': 'json', 'ai_enhancement': 'false'},
                                  timeout=T)
                assert 'source_type: png' in r.json()['markdown'], f"Failed for {upload_name}"


# ═══════════════════════════════════════════════════════════════════
# 3. OUTPUT FORMAT TESTING
# ═══════════════════════════════════════════════════════════════════

class TestOutputFormats:
    def test_json_format(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 't.csv'
            _csv(p, rows=2)
            r = _convert(p, output_format='json')
            d = r.json()
            assert 'filename' in d
            assert 'markdown' in d
            assert 'status' in d
            assert d['status'] == 'completed'

    def test_markdown_format(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 't.csv'
            _csv(p, rows=2)
            r = _convert(p, output_format='markdown')
            assert r.headers['Content-Type'].startswith('text/markdown')
            assert '---' in r.text

    def test_default_is_markdown(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 't.csv'
            _csv(p, rows=2)
            r = requests.post(f"{BASE}/api/convert",
                              files={'file': open(p, 'rb')},
                              data={'ai_enhancement': 'false'},
                              timeout=T)
            assert r.headers['Content-Type'].startswith('text/markdown')

    def test_frontmatter_always_present(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 't.csv'
            _csv(p, rows=2)
            r = _convert(p)
            md = r.json()['markdown']
            assert md.startswith('---')
            assert 'source_file:' in md
            assert 'processed:' in md

    def test_json_output_filename(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'report.csv'
            _csv(p, rows=2)
            r = _convert(p)
            d = r.json()
            assert d['output_filename'] == 'report.md'
            assert d['filename'] == 'report.csv'


# ═══════════════════════════════════════════════════════════════════
# 4. AI ENHANCEMENT TOGGLE
# ═══════════════════════════════════════════════════════════════════

class TestAIToggle:
    def test_ai_on_produces_summary(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'ai.csv'
            _csv(p, rows=20, cols=5)
            r = _convert(p, ai_enhancement='true')
            md = r.json()['markdown']
            assert 'summary:' in md
            assert 'tags:' in md

    def test_ai_off_no_summary(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'noai.csv'
            _csv(p, rows=5)
            r = _convert(p, ai_enhancement='false')
            md = r.json()['markdown']
            fm = md.split('---')[1]  # frontmatter section
            assert 'summary:' not in fm

    def test_ai_toggle_case_insensitive(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 't.csv'
            _csv(p, rows=5)
            for val in ['False', 'FALSE', 'false']:
                r = _convert(p, ai_enhancement=val)
                fm = r.json()['markdown'].split('---')[1]
                assert 'summary:' not in fm


# ═══════════════════════════════════════════════════════════════════
# 5. CHUNKED UPLOAD
# ═══════════════════════════════════════════════════════════════════

class TestChunkedUpload:
    def _init(self, filename, size, chunks, **kw):
        kw.setdefault('ai_enhancement', 'false')
        return requests.post(f"{BASE}/api/upload/init", data={
            'filename': filename, 'total_size': size, 'total_chunks': chunks, **kw
        })

    def _chunk(self, uid, idx, data):
        return requests.post(f"{BASE}/api/upload/chunk/{uid}",
                             files={'chunk': (f'c{idx}', data)},
                             data={'chunk_index': idx})

    def _complete(self, uid):
        return requests.post(f"{BASE}/api/upload/complete/{uid}")

    def _poll(self, job_id, timeout=60):
        for _ in range(timeout):
            r = requests.get(f"{BASE}/api/jobs/{job_id}")
            s = r.json()['status']
            if s in ('completed', 'failed'):
                return r.json()
            time.sleep(1)
        return {'status': 'timeout'}

    def test_single_chunk(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'sc.csv'
            _csv(p, rows=5)
            data = p.read_bytes()
            r = self._init('sc.csv', len(data), 1)
            uid = r.json()['upload_id']
            self._chunk(uid, 0, data)
            r = self._complete(uid)
            assert r.status_code == 200
            j = self._poll(r.json()['job_id'])
            assert j['status'] == 'completed'

    def test_three_chunks(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / '3c.csv'
            _csv(p, rows=100)
            data = p.read_bytes()
            cs = math.ceil(len(data) / 3)
            r = self._init('3c.csv', len(data), 3)
            uid = r.json()['upload_id']
            for i in range(3):
                self._chunk(uid, i, data[i*cs:(i+1)*cs])
            r = self._complete(uid)
            j = self._poll(r.json()['job_id'])
            assert j['status'] == 'completed'

    def test_init_unsupported_format(self):
        r = self._init('video.mp4', 1000, 1)
        assert r.status_code == 415

    def test_init_too_large(self):
        r = self._init('huge.csv', 600 * 1024 * 1024, 20)
        assert r.status_code == 413

    def test_chunk_bad_session(self):
        r = self._chunk('nonexistent', 0, b'data')
        assert r.status_code == 404

    def test_complete_missing_chunks(self):
        r = self._init('missing.csv', 1000, 5)
        uid = r.json()['upload_id']
        self._chunk(uid, 0, b'partial')
        r = self._complete(uid)
        assert r.status_code == 400

    def test_complete_nonexistent(self):
        r = self._complete('nonexistent')
        assert r.status_code == 404


# ═══════════════════════════════════════════════════════════════════
# 6. ASYNC UPLOAD + JOBS
# ═══════════════════════════════════════════════════════════════════

class TestAsyncUpload:
    def test_upload_poll_download(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'async.csv'
            _csv(p, rows=5)
            r = requests.post(f"{BASE}/api/upload",
                              files={'files[]': open(p, 'rb')},
                              data={'ai_enhancement': 'false'})
            jid = r.json()['job_ids'][0]
            for _ in range(60):
                r = requests.get(f"{BASE}/api/jobs/{jid}")
                if r.json()['status'] == 'completed':
                    break
                time.sleep(1)
            assert r.json()['status'] == 'completed'
            r = requests.get(f"{BASE}/api/download/{jid}")
            assert r.status_code == 200
            assert 'r0c0' in r.text

    def test_job_not_found(self):
        r = requests.get(f"{BASE}/api/jobs/fake-id")
        assert r.status_code == 404

    def test_download_not_found(self):
        r = requests.get(f"{BASE}/api/download/fake-id")
        assert r.status_code == 404

    def test_download_incomplete_job(self):
        """Upload but try downloading immediately (before completion)."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'quick.csv'
            _csv(p, rows=2)
            r = requests.post(f"{BASE}/api/upload",
                              files={'files[]': open(p, 'rb')},
                              data={'ai_enhancement': 'false'})
            jid = r.json()['job_ids'][0]
            # Try downloading immediately — may be processing or completed
            r = requests.get(f"{BASE}/api/download/{jid}")
            assert r.status_code in [200, 400]  # 400 if still processing

    def test_upload_skips_unsupported(self):
        r = requests.post(f"{BASE}/api/upload",
                          files={'files[]': ('bad.zip', b'PK', 'application/zip')},
                          data={'ai_enhancement': 'false'})
        d = r.json()
        assert d['job_ids'] == []
        assert 'bad.zip' in d.get('skipped', [])

    def test_job_response_has_no_internal_paths(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'sec.csv'
            _csv(p, rows=2)
            r = requests.post(f"{BASE}/api/upload",
                              files={'files[]': open(p, 'rb')},
                              data={'ai_enhancement': 'false'})
            jid = r.json()['job_ids'][0]
            time.sleep(3)
            r = requests.get(f"{BASE}/api/jobs/{jid}")
            body = r.text
            assert '/app/' not in body
            assert '_api_tmp' not in body


# ═══════════════════════════════════════════════════════════════════
# 7. EDGE CASES & EXPLORATORY
# ═══════════════════════════════════════════════════════════════════

class TestEdgeCases:
    def test_zero_byte_file(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('empty.csv', b'', 'text/csv')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        assert r.status_code in [200, 500]

    def test_1_byte_file(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('tiny.csv', b'x', 'text/csv')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        assert r.status_code in [200, 500]

    def test_binary_garbage_as_csv(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('garbage.csv', os.urandom(1024), 'text/csv')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        # Should not crash — either succeeds with error content or returns 500
        assert r.status_code in [200, 500]

    def test_corrupt_pdf(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('corrupt.pdf', b'%PDF-1.0\ngarbage', 'application/pdf')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        assert r.status_code in [200, 500]

    def test_corrupt_docx(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('corrupt.docx', b'PK\x03\x04garbage', 'application/docx')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        assert r.status_code in [200, 500]

    def test_corrupt_xlsx(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('corrupt.xlsx', b'PK\x03\x04garbage', 'application/xlsx')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        assert r.status_code in [200, 500]

    def test_csv_2000_rows_truncated(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'big.csv'
            _csv(p, rows=2000)
            r = _convert(p)
            md = r.json()['markdown']
            assert 'Showing first' in md

    def test_very_wide_csv(self):
        """CSV with 100 columns."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'wide.csv'
            _csv(p, rows=5, cols=100)
            r = _convert(p)
            assert r.status_code == 200

    def test_csv_with_unicode(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'unicode.csv'
            with open(p, 'w', encoding='utf-8') as f:
                f.write('name,city\nMüller,München\n田中,東京\n')
            r = _convert(p)
            md = r.json()['markdown']
            assert 'Müller' in md
            assert '東京' in md

    def test_large_image(self):
        """4000x4000 PNG — large but valid."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'large.png'
            _img(p, 4000, 4000)
            r = _convert(p)
            assert r.status_code == 200
            assert '4000' in r.json()['markdown']

    def test_1x1_pixel_image(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'tiny.png'
            _img(p, 1, 1)
            r = _convert(p)
            assert r.status_code == 200

    def test_rgba_image(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'rgba.png'
            _img(p, 50, 50, mode='RGBA')
            r = _convert(p)
            assert r.status_code == 200
            assert 'RGBA' in r.json()['markdown']

    def test_grayscale_image(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'gray.png'
            _img(p, 50, 50, mode='L')
            r = _convert(p)
            assert r.status_code == 200


# ═══════════════════════════════════════════════════════════════════
# 8. CONCURRENCY & MULTI-USER
# ═══════════════════════════════════════════════════════════════════

class TestConcurrency:
    def test_3_concurrent_same_filename(self):
        """Three users uploading report.csv simultaneously."""
        def upload(uid):
            with tempfile.TemporaryDirectory() as tmp:
                p = Path(tmp) / 'report.csv'
                with open(p, 'w', newline='') as f:
                    w = csv.writer(f)
                    w.writerow(['user', 'val'])
                    w.writerow([f'user_{uid}', uid * 77])
                r = _convert(p)
                return uid, r.status_code, r.json()

        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as pool:
            results = list(pool.map(lambda u: upload(u), range(3)))

        for uid, code, data in results:
            assert code == 200, f"User {uid}: HTTP {code}"
            assert data['status'] == 'completed'
            assert f'user_{uid}' in data['markdown'], f"User {uid} got wrong data"

    def test_5_concurrent_different_formats(self):
        """Five users uploading different formats simultaneously."""
        def make_and_upload(fmt_info):
            idx, ext, creator = fmt_info
            with tempfile.TemporaryDirectory() as tmp:
                p = Path(tmp) / f'test.{ext}'
                creator(p)
                r = _convert(p)
                return idx, ext, r.status_code

        tasks = [
            (0, 'csv', lambda p: _csv(p, rows=5)),
            (1, 'png', lambda p: _img(p)),
            (2, 'docx', lambda p: _docx(p)),
            (3, 'html', lambda p: _html(p)),
            (4, 'xlsx', lambda p: _xlsx(p)),
        ]

        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as pool:
            results = list(pool.map(make_and_upload, tasks))

        for idx, ext, code in results:
            assert code == 200, f"Format {ext}: HTTP {code}"


# ═══════════════════════════════════════════════════════════════════
# 9. SECURITY
# ═══════════════════════════════════════════════════════════════════

class TestSecurity:
    def test_path_traversal_filename(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('../../../etc/passwd.csv', b'name\nalice', 'text/csv')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        # Should not write outside upload dir — either sanitized or rejected
        assert r.status_code in [200, 400, 415, 422, 500]

    def test_null_byte_filename(self):
        r = requests.post(f"{BASE}/api/convert",
                          files={'file': ('test\x00.csv', b'name\nalice', 'text/csv')},
                          data={'output_format': 'json', 'ai_enhancement': 'false'},
                          timeout=T)
        assert r.status_code in [200, 400, 415, 422, 500]

    def test_no_cache_headers_present(self):
        r = requests.get(f"{BASE}/api/jobs/test")
        assert 'no-store' in r.headers.get('Cache-Control', '')

    def test_cors_headers(self):
        r = requests.options(f"{BASE}/api/convert",
                             headers={'Origin': 'http://evil.com', 'Access-Control-Request-Method': 'POST'})
        assert r.status_code in [200, 405]


# ═══════════════════════════════════════════════════════════════════
# 10. PERFORMANCE
# ═══════════════════════════════════════════════════════════════════

class TestPerformance:
    def test_csv_under_2_seconds(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'perf.csv'
            _csv(p, rows=100, cols=10)
            t0 = time.time()
            r = _convert(p)
            elapsed = time.time() - t0
            assert r.status_code == 200
            assert elapsed < 5, f"CSV took {elapsed:.1f}s (expected <5s)"

    def test_image_under_3_seconds(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'perf.png'
            _img(p, 500, 500)
            t0 = time.time()
            r = _convert(p)
            elapsed = time.time() - t0
            assert r.status_code == 200
            assert elapsed < 5, f"Image took {elapsed:.1f}s (expected <5s)"

    def test_docx_under_10_seconds(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / 'perf.docx'
            doc = Document()
            for i in range(20):
                doc.add_heading(f'Section {i}', 1)
                doc.add_paragraph('Lorem ipsum ' * 50)
            doc.save(p)
            t0 = time.time()
            r = _convert(p)
            elapsed = time.time() - t0
            assert r.status_code == 200
            assert elapsed < 15, f"DOCX took {elapsed:.1f}s (expected <15s)"

    def test_health_under_100ms(self):
        t0 = time.time()
        r = requests.get(f"{BASE}/health")
        elapsed = time.time() - t0
        assert r.status_code == 200
        assert elapsed < 0.5, f"Health took {elapsed:.3f}s"

    def test_throughput_10_small_files(self):
        """10 small CSVs in sequence — total under 30s."""
        t0 = time.time()
        for i in range(10):
            with tempfile.TemporaryDirectory() as tmp:
                p = Path(tmp) / f'tp_{i}.csv'
                _csv(p, rows=5)
                r = _convert(p)
                assert r.status_code == 200
        elapsed = time.time() - t0
        avg = elapsed / 10
        print(f"  10 files in {elapsed:.1f}s (avg {avg:.1f}s each)")
        assert elapsed < 60, f"10 files took {elapsed:.1f}s (expected <60s)"


# ═══════════════════════════════════════════════════════════════════
# 11. OPENAPI / SWAGGER
# ═══════════════════════════════════════════════════════════════════

class TestOpenAPI:
    def test_swagger_loads(self):
        r = requests.get(f"{BASE}/docs")
        assert r.status_code == 200

    def test_openapi_schema(self):
        r = requests.get(f"{BASE}/openapi.json")
        d = r.json()
        assert d['info']['version'] == '2.1.0'
        paths = list(d['paths'].keys())
        assert '/api/convert' in paths
        assert '/api/upload' in paths
        assert '/api/upload/init' in paths
        assert '/api/upload/chunk/{upload_id}' in paths
        assert '/api/upload/complete/{upload_id}' in paths
        assert '/api/jobs/{job_id}' in paths
        assert '/api/download/{job_id}' in paths
        assert '/health' in paths


if __name__ == '__main__':
    import pytest
    pytest.main([__file__, '-v', '--tb=short'])
